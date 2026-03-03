#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a Word document from Markdown content
"""

import docx
from docx.shared import Pt
from docx.oxml.ns import qn

def generate_word_manual(title, content):
    """Generate a Word document from Markdown content"""
    doc = docx.Document()

    # Add title
    heading = doc.add_heading(title, 0)
    heading.alignment = 1
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Add content
    for line in content.split('\n'):
        if line.startswith('#'):
            level = line.count('#')
            heading = doc.add_heading(line.strip('# '), level)
            for run in heading.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        elif line.startswith('- '):
            p = doc.add_paragraph()
            p.add_run('• ').font.name = 'Microsoft YaHei'
            p.add_run(line[2:]).font.name = 'Microsoft YaHei'
            for run in p.runs:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        elif line.startswith('|'):
            pass
        elif line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    return doc

if __name__ == '__main__':
    with open('/root/.openclaw/workspace/华宝通智能门锁锂电池说明书.md', 'r', encoding='utf-8') as f:
        content = f.read()

    doc = generate_word_manual('华宝通智能门锁锂电池说明书', content)
    doc.save('/root/.openclaw/workspace/华宝通智能门锁锂电池说明书.docx')
